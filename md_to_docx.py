"""
Convert COMPREHENSIVE_PROJECT_REPORT.md to DOCX with proper formatting.
Parses headings, paragraphs, tables, and lists.
"""
import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE

def parse_markdown_to_docx(md_path, docx_path):
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Table buffer
    table_buffer = []
    in_table = False
    table_headers = []

    def flush_table():
        """Write buffered table rows to document"""
        nonlocal table_buffer, table_headers
        if not table_buffer:
            return
        rows = [table_headers] + table_buffer
        cols = len(rows[0]) if rows else 0
        table = doc.add_table(rows=len(rows), cols=cols)
        table.style = 'Table Grid'
        for i, row in enumerate(rows):
            for j, cell_text in enumerate(row):
                cell = table.cell(i, j)
                cell.text = cell_text
                # Header row bold
                if i == 0:
                    for run in cell.paragraphs[0].runs:
                        run.bold = True
        table_buffer = []
        in_table = False

    # Process lines
    for line in lines:
        stripped = line.strip()

        # Empty line
        if not stripped:
            if in_table:
                flush_table()
            else:
                doc.add_paragraph()
            continue

        # Headings
        if stripped.startswith('# '):
            if in_table: flush_table()
            p = doc.add_heading(stripped[2:], level=1)
            continue
        if stripped.startswith('## '):
            if in_table: flush_table()
            p = doc.add_heading(stripped[3:], level=2)
            continue
        if stripped.startswith('### '):
            if in_table: flush_table()
            p = doc.add_heading(stripped[4:], level=3)
            continue

        # Table row
        if stripped.startswith('|'):
            # Remove leading/trailing |
            cells = [c.strip() for c in stripped.strip('|').split('|')]
            # Skip separator line (|---|---)
            if cells and all(re.fullmatch(r'-+', c) for c in cells if c):
                continue
            if not in_table:
                in_table = True
                table_headers = cells
            else:
                table_buffer.append(cells)
            continue

        # Regular paragraph (end table if inside)
        if in_table:
            flush_table()

        # Handle bold: **text** → format
        p = doc.add_paragraph()
        # Simple inline formatting (asterisks)
        # Split by ** pairs
        parts = re.split(r'(\*\*[^*]+\*\*)', stripped)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = p.add_run(part[2:-2])
                run.bold = True
            else:
                # Handle single asterisks for italic?
                if part.startswith('*') and part.endswith('*') and len(part) > 2:
                    run = p.add_run(part[1:-1])
                    run.italic = True
                else:
                    p.add_run(part)

    # Flush remaining table
    if in_table:
        flush_table()

    # Save
    doc.save(docx_path)
    print(f"✓ DOCX saved: {docx_path}")
    print(f"  Paragraphs: {len(doc.paragraphs)}")
    print(f"  Tables: {len(doc.tables)}")

# Run
parse_markdown_to_docx(
    'COMPREHENSIVE_PROJECT_REPORT.md',
    'COMPREHENSIVE_PROJECT_REPORT.docx'
)
