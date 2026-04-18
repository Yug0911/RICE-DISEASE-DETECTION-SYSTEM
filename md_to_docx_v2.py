"""
Convert COMPREHENSIVE_PROJECT_REPORT.md to DOCX with robust table handling
"""
import re
from docx import Document
from docx.shared import Pt

def parse_markdown_to_docx(md_path, docx_path):
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    doc = Document()
    # Default style
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    table_buffer = []
    in_table = False
    table_headers = []

    def flush_table():
        nonlocal table_buffer, table_headers
        if not table_buffer:
            return
        # Determine max columns
        all_rows = [table_headers] + table_buffer
        cols = max(len(row) for row in all_rows)
        rows = len(all_rows)
        table = doc.add_table(rows=rows, cols=cols)
        table.style = 'Table Grid'
        for i, row in enumerate(all_rows):
            for j in range(cols):
                cell_text = row[j] if j < len(row) else ''
                cell = table.cell(i, j)
                cell.text = cell_text
                if i == 0:  # header
                    for run in cell.paragraphs[0].runs:
                        run.bold = True
        table_buffer = []
        in_table = False

    for line in lines:
        stripped = line.strip()

        if not stripped:
            if in_table:
                flush_table()
            else:
                doc.add_paragraph()
            continue

        # Headings
        if stripped.startswith('# '):
            if in_table: flush_table()
            doc.add_heading(stripped[2:], level=1)
            continue
        if stripped.startswith('## '):
            if in_table: flush_table()
            doc.add_heading(stripped[3:], level=2)
            continue
        if stripped.startswith('### '):
            if in_table: flush_table()
            doc.add_heading(stripped[4:], level=3)
            continue

        # Table rows
        if stripped.startswith('|'):
            # Split and strip
            raw_cells = stripped.strip('|').split('|')
            cells = [c.strip() for c in raw_cells]
            # Check separator (e.g., |---|---|---|)
            if cells and all(re.fullmatch(r'-+', c) for c in cells if c):
                continue
            if not in_table:
                in_table = True
                table_headers = cells
            else:
                table_buffer.append(cells)
            continue

        # Regular text
        if in_table:
            flush_table()

        p = doc.add_paragraph()
        # Handle **bold**
        parts = re.split(r'(\*\*[^*]+\*\*)', stripped)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = p.add_run(part[2:-2])
                run.bold = True
            else:
                # Handle *italic*
                italic_match = re.fullmatch(r'\*([^*]+)\*', part)
                if italic_match:
                    run = p.add_run(italic_match.group(1))
                    run.italic = True
                else:
                    p.add_run(part)

    if in_table:
        flush_table()

    doc.save(docx_path)
    print(f"[OK] DOCX saved: {docx_path}")
    print(f"  Paragraphs: {len(doc.paragraphs)}")
    print(f"  Tables: {len(doc.tables)}")

parse_markdown_to_docx('COMPREHENSIVE_PROJECT_REPORT.md', 'COMPREHENSIVE_PROJECT_REPORT.docx')
